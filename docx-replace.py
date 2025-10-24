from docx import Document


def replace_in_paragraph(paragraph, old_text, new_text):
    """
    Заменяет все вхождения old_text на new_text в одном параграфе.
    Поведение имитирует Microsoft Word.
    """
    # Защита от пустого old_text или пустого параграфа
    if not old_text or not paragraph.runs:
        return

    # Объединяем текст всех run'ов параграфа в одну строку для поиска
    full_text = ''.join(run.text for run in paragraph.runs)

    # Если искомого текста нет — выходим
    if old_text not in full_text:
        return

    # Начинаем поиск с позиции 0
    start_search = 0

    # будем искать все вхождения по очереди
    while True:
        # Ищем позицию первого вхождения old_text, начиная с start_search
        pos = full_text.find(old_text, start_search)

        # Если больше нет вхождений — выходим из цикла
        if pos == -1:
            break

        # Конец заменяемого фрагмента (не включительно)
        end_pos = pos + len(old_text)
        # Получаем актуальный список run'ов (он может меняться при удалении)
        runs = paragraph.runs
        # current_offset — текущая позиция в full_text, соответствующая началу текущего run
        current_offset = 0
        # Список для хранения информации о run'ах, затронутых заменой:
        # каждый элемент: (индекс_run, начало_сегмента_в_run, конец_сегмента_в_run)
        affected_runs = []

        # Проходим по всем run'ам, чтобы понять, какие из них пересекаются с [pos, end_pos)
        for i, run in enumerate(runs):
            run_len = len(run.text)
            run_start = current_offset          # позиция начала этого run в full_text
            run_end = current_offset + run_len  # позиция конца этого run в full_text

            # Проверяем, пересекается ли этот run с заменяемым фрагментом
            if run_end <= pos or run_start >= end_pos:
                pass # Run полностью до или после заменяемого фрагмента — пропускаем
            else:
                # Этот run частично или полностью входит в заменяемый фрагмент

                # Где внутри run начинается заменяемый фрагмент?
                seg_start = max(0, pos - run_start)
                # Где внутри run заканчивается заменяемый фрагмент?
                seg_end = min(run_len, end_pos - run_start)
                # Сохраняем информацию об этом run
                affected_runs.append((i, seg_start, seg_end))
            # Сдвигаем offset на длину текущего run
            current_offset = run_end

        # Если по какой-то причине не нашли затронутые run — пропускаем это вхождение
        if not affected_runs:
            start_search = pos + 1
            continue

        # Определяем первый и последний затронутые run
        first_idx = affected_runs[0][0]   # индекс первого run, где начинается old_text
        last_idx = affected_runs[-1][0]   # индекс последнего run, где заканчивается old_text

        # === Часть 1: текст ДО old_text в первом run ===
        # Берём всё, что в первом run идёт до начала old_text
        before = runs[first_idx].text[:affected_runs[0][1]]
         # === Часть 2: текст ПОСЛЕ old_text в последнем run ===
        # Берём всё, что в последнем run идёт после конца old_text
        after = runs[last_idx].text[affected_runs[-1][2]:]

        # === Главное действие: замена текста ===
        # В ПЕРВОМ затронутом run оставляем "до" + новый текст
        runs[first_idx].text = before + new_text

        # Обновляем последний run, если он не первый
        if first_idx != last_idx:
            runs[last_idx].text = after

        # Удаляем промежуточные run (между first и last)
        for idx in range(last_idx - 1, first_idx, -1):
            paragraph._element.remove(runs[idx]._element)

        # Обновляем для следующей итерации
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        # Сдвигаем позицию поиска за пределы только что заменённого фрагмента
        start_search = pos + len(new_text)
        # Защита от бесконечного цикла, если old_text == new_text
        if old_text == new_text:
            start_search += 1


def docx_replace(doc, **kwargs):
    """Заменяет текст во всех элементах документа"""
    for old_text, new_text in kwargs.items():
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, old_text, new_text)
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for nested_paragraph in nested_cell.paragraphs:
                                    replace_in_paragraph(nested_paragraph, old_text, new_text)



if __name__ == "__main__":

    # Создаем документ с помощью docx
    doc = Document("document.docx")

    # производим замену в документе указав пары ключ значение
    docx_replace(doc, word1="word2", word3="word4")

    # производим замену в документе передав словарь
    my_dict = {"word5":"word6", "word7": "word8"}
    docx_replace(doc, **my_dict)


    # сохраняем полученный документ
    doc.save("replaced.docx")