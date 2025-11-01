# tests/test_docx_replace.py
import sys
import os
from pathlib import Path

import tempfile
# import pytest
from docx import Document
# from docx.text.run import Run
# from io import BytesIO

# Добавляем корень проекта в PYTHONPATH
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx_replace import docx_replace, replace_in_paragraph
# from docx.shared import Inches


# ------------------- Вспомогательные функции -------------------

def create_test_doc(paragraphs=None, tables=None, nested_tables=False):
    """Создаёт временный Document для тестов."""
    doc = Document()

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if tables:
        for table_data in tables:
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row in enumerate(table_data):
                for j, cell_text in enumerate(row):
                    table.cell(i, j).text = cell_text

            if nested_tables:
                # Добавим вложенную таблицу в первую ячейку
                nested = table.cell(0, 0).add_table(1, 1)
                nested.cell(0, 0).text = "nested князь"

    return doc


def get_all_text(doc):
    """Извлекает весь текст из документа для проверки."""
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            full_text.append(nested_cell.text)
    return "\n".join(full_text)


# ------------------- Тесты для replace_in_paragraph -------------------

class TestReplaceInParagraph:
    """Тесты для внутренней функции replace_in_paragraph."""

    def test_replace_single_run(self):
        """Замена внутри одного run."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Hello World!")
        replace_in_paragraph(p, "World", "Universe")
        assert p.text == "Hello Universe!"

    def test_replace_multiple_runs(self):
        """Замена текста, охватывающего несколько run'ов."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Hello ")
        p.add_run("World")
        p.add_run("!")
        replace_in_paragraph(p, "lo Wo", "XXX")
        assert p.text == "HelXXXrld!"

    def test_replace_multiple_occurrences(self):
        """Замена нескольких вхождений одного и того же текста."""
        doc = Document()
        p = doc.add_paragraph("cat and cat and cat")
        replace_in_paragraph(p, "cat", "dog")
        assert p.text == "dog and dog and dog"

    def test_replace_not_found(self):
        """Текст для замены не найден."""
        doc = Document()
        p = doc.add_paragraph("Hello World!")
        original_text = p.text
        replace_in_paragraph(p, "Python", "Java")
        assert p.text == original_text

    def test_replace_empty_old_text(self):
        """Заменяемый текст пуст."""
        doc = Document()
        p = doc.add_paragraph("Hello World!")
        original_text = p.text
        replace_in_paragraph(p, "", "Java")
        assert p.text == original_text

    def test_replace_empty_paragraph(self):
        """Параграф пуст или не содержит run'ов."""
        doc = Document()
        p = doc.add_paragraph("")
        replace_in_paragraph(p, "old", "new")
        assert p.text == ""

    def test_replace_same_text(self):
        """Заменяемый и новый текст совпадают."""
        doc = Document()
        p = doc.add_paragraph("Hello Hello World!")
        replace_in_paragraph(p, "Hello", "Hello")
        assert p.text == "Hello Hello World!"

    def test_replace_at_run_boundary(self):
        """Замена текста на границе двух run'ов."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Hello")
        p.add_run("World")
        replace_in_paragraph(p, "oW", "XX")
        assert p.text == "HellXXorld"

    def test_replace_at_paragraph_boundary(self):
        """Замена текста в начале и в конце параграфа."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Start")
        p.add_run("Middle")
        p.add_run("End")

        # Замена в начале
        replace_in_paragraph(p, "Star", "XXX")
        assert p.text == "XXXtMiddleEnd"

        # Восстановим и заменим в конце
        p.clear()
        p.add_run("Start")
        p.add_run("Middle")
        p.add_run("End")
        replace_in_paragraph(p, "End", "YYY")
        assert p.text == "StartMiddleYYY"

    def test_replace_with_empty_string(self):
        """Замена текста на пустую строку (удаление)."""
        doc = Document()
        p = doc.add_paragraph("Hello World!")
        replace_in_paragraph(p, " World", "")
        assert p.text == "Hello!"

    def test_replace_special_characters(self):
        """Замена текста с пробелами и специальными символами."""
        doc = Document()
        p = doc.add_paragraph("Line 1\nLine 2\tTab")
        replace_in_paragraph(p, "Line 2\t", "NewLine")
        assert p.text == "Line 1\nNewLineTab"


# ------------------- Тесты для docx_replace -------------------

class TestDocxReplace:
    """Тесты для основной функции docx_replace."""

    def test_simple_replacement(self):
        """Простая замена текста в параграфе."""
        doc = create_test_doc(paragraphs=["Hello князь world"])
        docx_replace(doc, князь="word6")
        assert "Hello word6 world" in get_all_text(doc)

    def test_multi_run_replacement(self):
        """Замена текста, охватывающего несколько run'ов."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Hel")
        p.add_run("lo князь!")
        docx_replace(doc, князь="word6")
        assert "Hello word6!" in get_all_text(doc)

    def test_multiple_occurrences(self):
        """Замена нескольких вхождений одного и того же текста."""
        doc = create_test_doc(paragraphs=["князь и князь снова князь"])
        docx_replace(doc, князь="word6")
        assert get_all_text(doc) == "word6 и word6 снова word6"

    def test_empty_old_text(self):
        """Проверка замены с пустым ключом."""
        doc = create_test_doc(paragraphs=["test"])
        docx_replace(doc, **{"": "ignored"})
        assert get_all_text(doc) == "test"

    def test_empty_new_text(self):
        """Замена текста на пустую строку (удаление)."""
        doc = create_test_doc(paragraphs=["Remove this князь now"])
        docx_replace(doc, князь="")
        assert get_all_text(doc) == "Remove this  now"

    def test_no_match(self):
        """Проверка случая, когда текст для замены не найден."""
        doc = create_test_doc(paragraphs=["Nothing to replace"])
        docx_replace(doc, князь="word6")
        assert get_all_text(doc) == "Nothing to replace"

    def test_same_old_and_new(self):
        """Заменяемый и новый текст совпадают."""
        doc = create_test_doc(paragraphs=["князь"])
        docx_replace(doc, князь="князь")
        assert get_all_text(doc) == "князь"

    def test_unicode_and_special_chars(self):
        """Замена текста с юникод символами и специальными знаками."""
        doc = create_test_doc(paragraphs=["Привет, князь Салтан!"])
        docx_replace(doc, князь="word6", Салтан="word8")
        assert "Привет, word6 word8!" in get_all_text(doc)

    def test_replacement_in_table(self):
        """Замена текста в таблице."""
        doc = create_test_doc(tables=[[["Cell with князь"]]])
        docx_replace(doc, князь="word6")
        assert "Cell with word6" in get_all_text(doc)

    def test_replacement_in_nested_table(self):
        """Замена текста во вложенной таблице."""
        doc = create_test_doc(tables=[[["Outer"]]], nested_tables=True)
        docx_replace(doc, князь="word6")
        text = get_all_text(doc)
        assert "nested word6" in text

    def test_overlapping_pattern(self):
        """Замена с перекрывающимися шаблонами."""
        doc = create_test_doc(paragraphs=["aaa"])
        docx_replace(doc, aa="X")
        assert get_all_text(doc) == "Xa"

    def test_multiple_keys(self):
        """Замена нескольких ключей одновременно."""
        doc = create_test_doc(paragraphs=["князь Салтан сказал: ты, куда?"])
        docx_replace(doc, **{"князь": "word6", "Салтан": "word8", "ты, куда": "word9"})
        assert "word6 word8 сказал: word9?" in get_all_text(doc)

    def test_replace_in_paragraphs_only(self):
        """Замена текста только в параграфах."""
        doc = Document()
        doc.add_paragraph("Original text here.")
        doc.add_paragraph("Another line with Original.")
        docx_replace(doc, Original="New", text="word")
        assert doc.paragraphs[0].text == "New word here."
        assert doc.paragraphs[1].text == "Another line with New."

    def test_replace_in_table_cells(self):
        """Замена текста в ячейках таблицы."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        cell1 = table.cell(0, 0)
        cell2 = table.cell(0, 1)
        cell1.text = "Cell 1 text Original"
        cell2.text = "Cell 2 Original text"
        docx_replace(doc, Original="Replaced", text="X")
        assert table.cell(0, 0).text == "Cell 1 X Replaced"
        assert table.cell(0, 1).text == "Cell 2 Replaced X"

    def test_replace_in_nested_elements_structure(self):
        """Замена текста с изменением структуры вложенных элементов."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("AA")
        p.add_run("BB")
        p.add_run("CC")
        docx_replace(doc, AABB="X")
        runs_after = p.runs
        assert len(runs_after) == 3
        assert runs_after[0].text == "X"
        assert runs_after[1].text == ""
        assert runs_after[2].text == "CC"

    def test_replace_multiple_keys_in_one_call(self):
        """Замена нескольких ключей в одном вызове функции."""
        doc = Document()
        p = doc.add_paragraph("First Second Third")
        docx_replace(doc, First="1", Second="2", Third="3")
        assert p.text == "1 2 3"

    def test_replace_empty_dict(self):
        """Проверка работы с пустым словарем замен."""
        doc = Document()
        p = doc.add_paragraph("Hello World!")
        original_text = p.text
        docx_replace(doc)
        assert p.text == original_text

    def test_replace_no_paragraphs_no_tables(self):
        """Проверка работы с пустым документом."""
        doc = Document()
        docx_replace(doc, test="replacement")
        assert len(doc.paragraphs) == 0
        assert len(doc.tables) == 0

    def test_replace_with_runs_merging_logic(self):
        """Проверка логики объединения run'ов при замене."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Start ")
        p.add_run("Middle")
        p.add_run(" End")
        docx_replace(doc, **{"rt Mid": "XX"})
        expected_text = "StaXXdle End"
        assert p.text == expected_text
        assert len(p.runs) == 3

    def test_run_structure_preserved_after_replacement(self):
        """Проверка сохранения структуры документа после замены."""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Start ")
        p.add_run("кня")
        p.add_run("зь end")
        docx_replace(doc, князь="word6")

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir) / "test.docx"
            doc.save(tmp_path)
            assert tmp_path.exists()


# ------------------- Запуск через pytest -------------------

if __name__ == "__main__":
    # Пример запуска: python -m pytest tests/test_docx_replace.py -v
    pass