import sys
import os

import re
from collections import Counter

# Добавляем корень проекта в PYTHONPATH
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx_replace import docx_replace

document1 = "Pushkin_AS.docx"
document2 = "replaced.docx"


# Создаем документ с помощью docx
doc = Document(document1)
# производим замену в документе указав пары ключ значение
# docx_replace(doc, Лебедь="word3", гости="word4")
# производим замену в документе передав словарь
my_dict = {"князь":"word6", "Салтан": "word8", "ты, куда": "word9"}
docx_replace(doc, **my_dict)
# сохраняем полученный документ
doc.save(document2)


def count_words_in_docx(file_path):
    """
    Читает .docx документ и возвращает словарь с количеством повторений каждого слова.

    :param file_path: путь к .docx файлу
    :return: словарь {слово: количество}
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f"Не удалось открыть документ: {e}")

    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Объединяем весь текст и приводим к нижнему регистру
    text = ' '.join(full_text)

    # Извлекаем только слова (без знаков препинания и цифр)
    # words = re.findall(r'\b\w+\b', text)
    words = [word for word in re.findall(r'\b\w+\b', text) if len(word) > 2]

    # Считаем частоту слов
    word_counts = Counter(words)

    # Сортируем по убыванию и создаём упорядоченный словарь
    sorted_word_counts = dict(sorted(word_counts.items(), key=lambda item: item[1], reverse=True))

    return sorted_word_counts

def compare_dicts(dict1, dict2):
    """
    Сравнивает два словаря и возвращает два новых словаря,
    содержащих только отличающиеся пары (ключ: значение).

    :param dict1: первый словарь
    :param dict2: второй словарь
    :return: (diff1, diff2) — кортеж из двух словарей
    """
    all_keys = set(dict1.keys()) | set(dict2.keys())

    diff1 = {}
    diff2 = {}

    for key in all_keys:
        val1 = dict1.get(key)
        val2 = dict2.get(key)

        # Включаем ключ, если значения отличаются (включая случай, когда одного нет)
        if val1 != val2:
            if key in dict1:
                diff1[key] = val1
            if key in dict2:
                diff2[key] = val2

    return diff1, diff2

def compare_dicts(dict1, dict2):
    """
    Сравнивает два словаря и возвращает два новых словаря,
    содержащих только отличающиеся пары (ключ: значение).

    :param dict1: первый словарь
    :param dict2: второй словарь
    :return: (diff1, diff2) — кортеж из двух словарей
    """
    all_keys = set(dict1.keys()) | set(dict2.keys())

    diff1 = {}
    diff2 = {}

    for key in all_keys:
        val1 = dict1.get(key)
        val2 = dict2.get(key)

        # Включаем ключ, если значения отличаются (включая случай, когда одного нет)
        if val1 != val2:
            if key in dict1:
                diff1[key] = val1
            if key in dict2:
                diff2[key] = val2

    return diff1, diff2


def count_substrings_in_docx(docx_path, mapping):
    """
    Считает, сколько раз каждая строка из ключей и значений словаря `mapping`
    встречается как подстрока в тексте документа .docx.

    Аргументы:
        docx_path (str): Путь к файлу .docx.
        mapping (dict): Словарь {ключ: значение}, где ключ и значение — строки.

    Возвращает:
        dict: Словарь вида {строка: количество_вхождений, ...}
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise ValueError(f"Не удалось открыть документ: {e}")

    # Собираем весь текст из параграфов
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    # Формируем множество всех строк для поиска (ключи + значения), приведённых к str
    search_strings = {str(k) for k in mapping.keys()} | {str(v) for v in mapping.values()}

    # Считаем вхождения каждой строки
    result = {}
    for s in search_strings:
        result[s] = full_text.count(s)

    return result

result1 = count_words_in_docx(document1)
result2 = count_words_in_docx(document2)
print(f'10 самых повторяющихся слов в документе {document1}:')
print(dict(list(result1.items())[:10]), '\n')

print('Словарь сров с заменами:')
print(my_dict, '\n')

diff1, diff2 = compare_dicts(result1, result2)
print(f'Изменения в документе {document1}, относительно {document2}:')
print(diff1)
print(f'Изменения в документе {document2}, относительно {document1}:')
print(diff2,'\n')

print('Показывает, сколько раз каждая строка из ключей и значений словаря встречается как подстрока в тексте документов')
result = count_substrings_in_docx(document1, my_dict)
print(document1, result)
result = count_substrings_in_docx(document2, my_dict)
print(document2, result)
